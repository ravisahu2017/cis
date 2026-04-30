import fitz  # PyMuPDF
import pytesseract
from docx import Document
from PIL import Image

from tools.logger import logger

def read_pdf(file_path: str) -> str:
    text_content = ""
    doc = fitz.open(file_path)
    for page_num in range(len(doc)):
        page = doc.load_page(page_num)
        text_content += page.get_text()
    doc.close()
    logger.info(f"Extracted {len(text_content)} characters from {file_path}")
    return text_content

def read_docx(file_path: str) -> str:
    """Extract text content from DOCX files"""
    try:
        doc = Document(file_path)
        text_content = ""
        
        # Extract text from paragraphs
        for paragraph in doc.paragraphs:
            text_content += paragraph.text + "\n"
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text_content += cell.text + "\t"
                text_content += "\n"
        
        logger.info(f"Extracted {len(text_content)} characters from DOCX file")
        return text_content.strip()
        
    except Exception as e:
        logger.error(f"Error reading DOCX file: {str(e)}")
        raise e

def read_txt(file_path: str) -> str:
    """Extract text content from TXT files"""
    try:
        with open(file_path, 'r', encoding='utf-8') as file:
            text_content = file.read()
        logger.info(f"Extracted {len(text_content)} characters from TXT file")
        return text_content
    except Exception as e:
        logger.error(f"Error reading TXT file: {str(e)}")
        raise e

def read_image(file_path: str) -> str:
    """Extract text content from image files"""
    try:
        image = Image.open(file_path)
        text_content = pytesseract.image_to_string(image)
        logger.info(f"Extracted {len(text_content)} characters from image file")
        return text_content
    except Exception as e:
        logger.error(f"Error reading image file: {str(e)}")
        raise e

